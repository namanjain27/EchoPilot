"""
Document processing and ingestion pipeline for EchoPilot
Handles document preprocessing, chunking, and metadata extraction
"""

import logging
import os
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import hashlib
import mimetypes
from dataclasses import dataclass
from datetime import datetime
import json
import time

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available - PDF processing disabled")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - Word document processing disabled")

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("pytesseract or Pillow not available - OCR processing disabled")


@dataclass
class DocumentChunk:
    """Represents a processed document chunk"""
    content: str
    chunk_id: str
    source_file: str
    chunk_index: int
    metadata: Dict[str, Any]


class DocumentProcessor:
    """Processes documents for ingestion into knowledge base"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor
        
        Args:
            chunk_size: Maximum size of each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Supported file types
        self.supported_types = {
            '.txt': self._process_text_file,
            '.md': self._process_text_file,
            '.json': self._process_json_file,
            '.csv': self._process_csv_file
        }
        
        # Add PDF support if available
        if PDF_AVAILABLE:
            self.supported_types['.pdf'] = self._process_pdf_file
        
        # Add Word document support if available
        if DOCX_AVAILABLE:
            self.supported_types['.docx'] = self._process_docx_file
            self.supported_types['.doc'] = self._process_docx_file
        
        # Add image OCR support if available
        if OCR_AVAILABLE:
            self.supported_types.update({
                '.jpg': self._process_image_file,
                '.jpeg': self._process_image_file,
                '.png': self._process_image_file,
                '.bmp': self._process_image_file,
                '.tiff': self._process_image_file
            })
        
        logger.info(f"Document processor initialized with chunk_size={chunk_size}, overlap={chunk_overlap}")
    
    def process_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Process a single file into chunks
        
        Args:
            file_path: Path to the file
            metadata: Optional metadata to attach to chunks
            
        Returns:
            List of document chunks
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return []
            
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            # Process file based on type
            processor = self.supported_types[file_extension]
            text_content = processor(str(file_path))
            
            if not text_content:
                logger.warning(f"No content extracted from: {file_path}")
                return []
            
            # Generate base metadata
            base_metadata = self._generate_file_metadata(str(file_path))
            if metadata:
                base_metadata.update(metadata)
            
            # Split into chunks
            chunks = self._split_text_into_chunks(text_content, str(file_path), base_metadata)
            
            logger.info(f"Processed {file_path}: {len(chunks)} chunks generated")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return []
    
    def process_directory(self, directory_path: str, 
                         recursive: bool = True,
                         metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Process all supported files in a directory
        
        Args:
            directory_path: Path to directory
            recursive: Whether to process subdirectories
            metadata: Optional metadata to attach to all chunks
            
        Returns:
            List of all document chunks
        """
        try:
            directory = Path(directory_path)
            if not directory.exists() or not directory.is_dir():
                logger.error(f"Directory not found or invalid: {directory_path}")
                return []
            
            all_chunks = []
            
            # Get all files
            if recursive:
                files = directory.rglob('*')
            else:
                files = directory.glob('*')
            
            # Process each supported file
            for file_path in files:
                if file_path.is_file() and file_path.suffix.lower() in self.supported_types:
                    file_chunks = self.process_file(str(file_path), metadata)
                    all_chunks.extend(file_chunks)
            
            logger.info(f"Processed directory {directory_path}: {len(all_chunks)} total chunks")
            return all_chunks
            
        except Exception as e:
            logger.error(f"Error processing directory {directory_path}: {e}")
            return []
    
    def process_text_content(self, content: str, source_name: str, 
                           metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Process raw text content into chunks
        
        Args:
            content: Text content to process
            source_name: Name/identifier for the content source
            metadata: Optional metadata
            
        Returns:
            List of document chunks
        """
        try:
            # Generate base metadata
            base_metadata = {
                'source': source_name,
                'type': 'text_content',
                'processed_at': datetime.now().isoformat(),
                'content_length': len(content)
            }
            
            if metadata:
                base_metadata.update(metadata)
            
            # Split into chunks
            chunks = self._split_text_into_chunks(content, source_name, base_metadata)
            
            logger.info(f"Processed text content '{source_name}': {len(chunks)} chunks generated")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing text content '{source_name}': {e}")
            return []
    
    def _process_text_file(self, file_path: str) -> str:
        """Process plain text or markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""
    
    def _process_json_file(self, file_path: str) -> str:
        """Process JSON file by converting to readable text"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text format
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error reading JSON file {file_path}: {e}")
            return ""
    
    def _process_csv_file(self, file_path: str) -> str:
        """Process CSV file by converting to text"""
        try:
            import csv
            import io
            
            text_content = []
            with open(file_path, 'r', encoding='utf-8') as f:
                csv_reader = csv.reader(f)
                for row in csv_reader:
                    text_content.append(' | '.join(row))
            
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error reading CSV file {file_path}: {e}")
            return ""
    
    def _process_pdf_file(self, file_path: str) -> str:
        """Process PDF file by extracting text"""
        if not PDF_AVAILABLE:
            logger.error("PDF processing not available - install PyPDF2")
            return ""

        try:
            text_content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1} in {file_path}: {e}")
                        continue
            
            if not text_content:
                logger.warning(f"No text content extracted from PDF: {file_path}")
                return ""
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            return ""
    
    def _process_docx_file(self, file_path: str) -> str:
        """Process Word document by extracting text"""
        if not DOCX_AVAILABLE:
            logger.error("Word document processing not available - install python-docx")
            return ""
        
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                logger.warning(f"No text content extracted from Word document: {file_path}")
                return ""
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading Word document {file_path}: {e}")
            return ""
    
    def _process_image_file(self, file_path: str) -> str:
        """Process image file using OCR to extract text"""
        if not OCR_AVAILABLE:
            logger.error("OCR processing not available - install pytesseract and Pillow")
            return ""
        
        try:
            # Open and process image
            with Image.open(file_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Perform OCR
                extracted_text = pytesseract.image_to_string(img)
                
                if not extracted_text.strip():
                    logger.warning(f"No text extracted from image: {file_path}")
                    return ""
                
                # Add image metadata to extracted text
                metadata_text = f"--- Text extracted from image: {os.path.basename(file_path)} ---\n"
                return metadata_text + extracted_text.strip()
                
        except Exception as e:
            logger.error(f"Error processing image file {file_path}: {e}")
            return ""
    
    def _split_text_into_chunks(self, text: str, source: str, 
                               metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into overlapping chunks with robust infinite loop protection"""
        chunks = []
        
        if len(text) <= self.chunk_size:
            # Text is small enough to fit in one chunk
            chunk_id = self._generate_chunk_id(source, 0)
            chunk = DocumentChunk(
                content=text,
                chunk_id=chunk_id,
                source_file=source,
                chunk_index=0,
                metadata={**metadata, 'chunk_size': len(text)}
            )
            chunks.append(chunk)
            return chunks
        
        # Split into overlapping chunks with safety limits
        start = 0
        chunk_index = 0
        max_iterations = (len(text) // (self.chunk_size - self.chunk_overlap)) + 10  # Safety buffer
        iteration_count = 0
        
        while start < len(text) and iteration_count < max_iterations:
            iteration_count += 1
            
            # Determine chunk end position
            end = start + self.chunk_size
            
            if end >= len(text):
                # Last chunk
                chunk_content = text[start:]
            else:
                # Find a good breaking point (end of sentence or paragraph)
                chunk_content = text[start:end]
                
                # Try to break at sentence end
                last_period = chunk_content.rfind('.')
                last_newline = chunk_content.rfind('\n')
                
                if last_period > len(chunk_content) * 0.8:
                    chunk_content = text[start:start + last_period + 1]
                elif last_newline > len(chunk_content) * 0.8:
                    chunk_content = text[start:start + last_newline + 1]
                # Otherwise use the full chunk
            
            # Skip empty chunks
            if not chunk_content.strip():
                start += max(1, self.chunk_size // 2)  # Move forward to avoid infinite loop
                continue
            
            # Create chunk
            chunk_id = self._generate_chunk_id(source, chunk_index)
            chunk = DocumentChunk(
                content=chunk_content.strip(),
                chunk_id=chunk_id,
                source_file=source,
                chunk_index=chunk_index,
                metadata={**metadata, 'chunk_size': len(chunk_content)}
            )
            chunks.append(chunk)
            
            # Calculate next start position with safety checks
            current_chunk_length = len(chunk_content)
            next_start = start + current_chunk_length - self.chunk_overlap
            
            # Ensure we're making progress
            if next_start <= start:
                # Force progress if we're not advancing
                next_start = start + max(1, current_chunk_length // 2)
                logger.warning(f"Forced progress in chunking for {source} at position {start}")
            
            start = next_start
            chunk_index += 1
            
            # Additional safety check for very large files
            if chunk_index > 10000:  # Arbitrary large limit
                logger.warning(f"Chunking stopped at {chunk_index} chunks for {source} - may be too large")
                break
        
        if iteration_count >= max_iterations:
            logger.error(f"Chunking iteration limit reached for {source} - stopping to prevent infinite loop")
        
        logger.info(f"Created {len(chunks)} chunks from {source} in {iteration_count} iterations")
        return chunks
    
    def _generate_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Generate metadata for a file"""
        try:
            path = Path(file_path)
            stat = path.stat()
            
            return {
                'source': str(path),
                'filename': path.name,
                'file_extension': path.suffix.lower(),
                'file_size': stat.st_size,
                'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'processed_at': datetime.now().isoformat(),
                'mime_type': mimetypes.guess_type(str(path))[0] or 'unknown'
            }
        except Exception as e:
            logger.error(f"Error generating metadata for {file_path}: {e}")
            return {'source': file_path, 'processed_at': datetime.now().isoformat()}
    
    def _generate_chunk_id(self, source: str, chunk_index: int) -> str:
        """Generate unique ID for a chunk"""
        content_hash = hashlib.md5(f"{source}_{chunk_index}".encode()).hexdigest()[:8]
        return f"{Path(source).stem}_{chunk_index}_{content_hash}"
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get statistics about the processor configuration"""
        return {
            'chunk_size': self.chunk_size,
            'chunk_overlap': self.chunk_overlap,
            'supported_extensions': list(self.supported_types.keys()),
            'processor_version': '2.0.0',
            'features': {
                'pdf_support': PDF_AVAILABLE,
                'docx_support': DOCX_AVAILABLE,
                'ocr_support': OCR_AVAILABLE
            }
        }
    
    def validate_file_type(self, file_path: str) -> Tuple[bool, str]:
        """Validate if file type is supported
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (is_supported, message)
        """
        try:
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            if file_extension in self.supported_types:
                return True, f"File type {file_extension} is supported"
            else:
                supported_types = ', '.join(self.supported_types.keys())
                return False, f"File type {file_extension} not supported. Supported types: {supported_types}"
                
        except Exception as e:
            return False, f"Error validating file: {e}"
    
    def get_file_size_mb(self, file_path: str) -> float:
        """Get file size in MB
        
        Args:
            file_path: Path to the file
            
        Returns:
            File size in MB
        """
        try:
            size_bytes = os.path.getsize(file_path)
            return size_bytes / (1024 * 1024)
        except Exception as e:
            logger.error(f"Error getting file size for {file_path}: {e}")
            return 0.0

class IngestionPipeline:
    """Complete document ingestion pipeline"""
    
    def __init__(self, knowledge_base_manager, document_processor: Optional[DocumentProcessor] = None):
        """
        Initialize ingestion pipeline
        
        Args:
            knowledge_base_manager: KnowledgeBaseManager instance
            document_processor: Optional DocumentProcessor instance
        """
        self.kb_manager = knowledge_base_manager
        self.processor = document_processor or DocumentProcessor()
        
        logger.info("Document ingestion pipeline initialized")
    
    def ingest_file(self, file_path: str, kb_type: str, 
                   metadata: Optional[Dict[str, Any]] = None) -> bool:
        """
        Ingest a single file into knowledge base with timeout protection
        
        Args:
            file_path: Path to file
            kb_type: Knowledge base type ('internal' or 'general')
            metadata: Optional additional metadata
            
        Returns:
            True if successful, False otherwise
        """
        start_time = time.time()
        
        try:
            from .knowledge_base import KnowledgeBaseType
            
            # Convert string to enum
            kb_enum = KnowledgeBaseType.INTERNAL if kb_type.lower() == 'internal' else KnowledgeBaseType.GENERAL
            
            logger.info(f"IngestionPipeline.ingest_file: Processing '{file_path}' with kb_type='{kb_type}' -> enum={kb_enum.value}")
            chunks = self.processor.process_file(file_path, metadata)
            
            # Check if processing took too long (basic timeout protection)
            processing_time = time.time() - start_time
            if processing_time > 300:  # 5 minutes timeout
                logger.error(f"File processing took too long ({processing_time:.2f}s) for {file_path}")
                return False
            
            if not chunks:
                logger.warning(f"No chunks generated from file: {file_path}")
                return False
            
            logger.info(f"Generated {len(chunks)} chunks, adding to knowledge base")
            
            # Prepare for knowledge base ingestion
            documents = [chunk.content for chunk in chunks]
            metadatas = [chunk.metadata for chunk in chunks]
            ids = [chunk.chunk_id for chunk in chunks]
            
            # Add to knowledge base
            logger.info(f"Adding {len(documents)} documents to knowledge base with kb_type={kb_enum.value}")
            success = self.kb_manager.add_documents(
                documents=documents,
                metadatas=metadatas,
                kb_type=kb_enum,
                ids=ids
            )
            
            processing_time = time.time() - start_time
            
            if success:
                logger.info(f"Successfully ingested {len(chunks)} chunks from {file_path} in {processing_time:.2f}s")
            else:
                logger.error(f"Failed to add chunks to knowledge base for {file_path}")
            
            return success
            
        except Exception as e:
            logger.error(f"Error ingesting file {file_path}: {e}")
            return False
    
    def ingest_directory(self, directory_path: str, kb_type: str,
                        recursive: bool = True,
                        metadata: Optional[Dict[str, Any]] = None) -> Tuple[int, int]:
        """
        Ingest all files in a directory
        
        Args:
            directory_path: Path to directory
            kb_type: Knowledge base type
            recursive: Process subdirectories
            metadata: Optional metadata for all files
            
        Returns:
            Tuple of (successful_files, total_files)
        """
        try:
            # Process all files in directory
            all_chunks = self.processor.process_directory(directory_path, recursive, metadata)
            
            if not all_chunks:
                logger.warning(f"No chunks generated from directory: {directory_path}")
                return 0, 0
            
            from .knowledge_base import KnowledgeBaseType
            kb_enum = KnowledgeBaseType.INTERNAL if kb_type.lower() == 'internal' else KnowledgeBaseType.GENERAL
            
            # Group chunks by source file
            files_processed = {}
            for chunk in all_chunks:
                if chunk.source_file not in files_processed:
                    files_processed[chunk.source_file] = []
                files_processed[chunk.source_file].append(chunk)
            
            successful_files = 0
            total_files = len(files_processed)
            
            # Ingest each file's chunks
            for source_file, chunks in files_processed.items():
                try:
                    documents = [chunk.content for chunk in chunks]
                    metadatas = [chunk.metadata for chunk in chunks]
                    ids = [chunk.chunk_id for chunk in chunks]
                    
                    success = self.kb_manager.add_documents(
                        documents=documents,
                        metadatas=metadatas,
                        kb_type=kb_enum,
                        ids=ids
                    )
                    
                    if success:
                        successful_files += 1
                        logger.info(f"Ingested {len(chunks)} chunks from {source_file}")
                    
                except Exception as e:
                    logger.error(f"Error ingesting chunks from {source_file}: {e}")
            
            logger.info(f"Directory ingestion complete: {successful_files}/{total_files} files successful")
            return successful_files, total_files
            
        except Exception as e:
            logger.error(f"Error ingesting directory {directory_path}: {e}")
            return 0, 0
    
    def ingest_text_batch(self, text_items: List[Dict[str, Any]], kb_type: str) -> bool:
        """
        Ingest a batch of text items
        
        Args:
            text_items: List of dicts with 'content', 'source', and optional 'metadata'
            kb_type: Knowledge base type
            
        Returns:
            True if successful, False otherwise
        """
        try:
            from .knowledge_base import KnowledgeBaseType
            kb_enum = KnowledgeBaseType.INTERNAL if kb_type.lower() == 'internal' else KnowledgeBaseType.GENERAL
            
            all_chunks = []
            
            # Process each text item
            for item in text_items:
                content = item.get('content', '')
                source = item.get('source', 'unknown')
                metadata = item.get('metadata', {})
                
                chunks = self.processor.process_text_content(content, source, metadata)
                all_chunks.extend(chunks)
            
            if not all_chunks:
                logger.warning("No chunks generated from text batch")
                return False
            
            # Prepare for knowledge base
            documents = [chunk.content for chunk in all_chunks]
            metadatas = [chunk.metadata for chunk in all_chunks]
            ids = [chunk.chunk_id for chunk in all_chunks]
            
            # Add to knowledge base
            success = self.kb_manager.add_documents(
                documents=documents,
                metadatas=metadatas,
                kb_type=kb_enum,
                ids=ids
            )
            
            if success:
                logger.info(f"Successfully ingested {len(all_chunks)} chunks from text batch")
            
            return success
            
        except Exception as e:
            logger.error(f"Error ingesting text batch: {e}")
            return False