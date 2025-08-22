#!/usr/bin/env python3
"""
Amazon Officers and Directors Data Scraper
Scrapes officer and director information from Amazon's IR page and exports to DOCX format.
"""

import requests
from bs4 import BeautifulSoup
import time
import re
from docx import Document
from docx.shared import Inches
from typing import List, Dict, Optional
import sys


class AmazonOfficersScraper:
    def __init__(self):
        self.base_url = "https://ir.aboutamazon.com/officers-and-directors/default.aspx"
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        })
        self.officers_data = []

    def fetch_page(self) -> Optional[BeautifulSoup]:
        """Fetch the officers and directors page with retry logic."""
        max_retries = 3
        for attempt in range(max_retries):
            try:
                print(f"Attempting to fetch page (attempt {attempt + 1}/{max_retries})...")
                response = self.session.get(self.base_url, timeout=30)
                
                if response.status_code == 200:
                    print("✓ Successfully fetched the page")
                    return BeautifulSoup(response.content, 'html.parser')
                elif response.status_code == 403:
                    print(f"Access denied (403). Trying with different headers...")
                    # Try with minimal headers
                    self.session.headers.update({
                        'User-Agent': 'Mozilla/5.0 (compatible; Research Bot)',
                    })
                    time.sleep(2)
                else:
                    print(f"HTTP {response.status_code}: {response.reason}")
                    
            except requests.exceptions.RequestException as e:
                print(f"Request failed: {e}")
                if attempt < max_retries - 1:
                    time.sleep(5)
                    
        return None

    def extract_officer_data(self, soup: BeautifulSoup) -> List[Dict]:
        """Extract officer and director data from the parsed HTML."""
        officers = []
        
        # Multiple possible selectors for officer information
        selectors_to_try = [
            '.officer-profile',
            '.director-profile', 
            '.executive-profile',
            '[class*="officer"]',
            '[class*="director"]',
            '[class*="executive"]',
            '.bio-section',
            '.leadership-profile'
        ]
        
        print("Analyzing page structure...")
        
        # Try different selectors
        found_profiles = False
        for selector in selectors_to_try:
            profiles = soup.select(selector)
            if profiles:
                print(f"Found {len(profiles)} profiles using selector: {selector}")
                found_profiles = True
                break
        
        if not found_profiles:
            # Fallback: look for common patterns in text
            print("No standard selectors found. Analyzing page content...")
            text_content = soup.get_text()
            
            # Look for typical executive titles
            executive_patterns = [
                r'Chief Executive Officer|CEO',
                r'Chief Financial Officer|CFO', 
                r'Chief Operating Officer|COO',
                r'President',
                r'Vice President|VP',
                r'Director',
                r'Chairman',
                r'Board Member'
            ]
            
            # Try to find sections with these patterns
            for pattern in executive_patterns:
                matches = re.finditer(pattern, text_content, re.IGNORECASE)
                for match in matches:
                    print(f"Found potential executive mention: {match.group()}")
        
        # Generic extraction approach
        # Look for div or section tags that might contain officer info
        potential_sections = soup.find_all(['div', 'section', 'article'], 
                                         class_=re.compile(r'(officer|director|executive|leadership|bio)', re.I))
        
        if potential_sections:
            print(f"Found {len(potential_sections)} potential officer sections")
            
            for section in potential_sections:
                officer_data = self.parse_officer_section(section)
                if officer_data:
                    officers.append(officer_data)
        
        # If still no structured data found, try to extract from tables
        tables = soup.find_all('table')
        for table in tables:
            table_data = self.parse_officer_table(table)
            officers.extend(table_data)
        
        return officers

    def parse_officer_section(self, section) -> Optional[Dict]:
        """Parse an individual officer section."""
        officer = {}
        
        # Try to find name
        name_selectors = ['h1', 'h2', 'h3', 'h4', '.name', '.officer-name', '.director-name']
        for selector in name_selectors:
            name_elem = section.select_one(selector)
            if name_elem:
                officer['name'] = name_elem.get_text().strip()
                break
        
        # Try to find position/title
        title_selectors = ['.title', '.position', '.role', 'h4', 'h5', '.officer-title']
        for selector in title_selectors:
            title_elem = section.select_one(selector)
            if title_elem and title_elem.get_text().strip() != officer.get('name', ''):
                officer['position'] = title_elem.get_text().strip()
                break
        
        # Try to find description/bio
        bio_selectors = ['.bio', '.description', '.officer-bio', 'p']
        bio_text = []
        for selector in bio_selectors:
            bio_elems = section.select(selector)
            for elem in bio_elems:
                text = elem.get_text().strip()
                if len(text) > 50:  # Likely a biography
                    bio_text.append(text)
        
        if bio_text:
            officer['description'] = ' '.join(bio_text)
        
        # Try to find committee membership
        committee_keywords = ['committee', 'board', 'audit', 'compensation', 'nominating', 'governance']
        committee_info = []
        
        all_text = section.get_text().lower()
        for keyword in committee_keywords:
            if keyword in all_text:
                # Extract sentences containing committee information
                sentences = section.get_text().split('.')
                for sentence in sentences:
                    if keyword in sentence.lower():
                        committee_info.append(sentence.strip())
        
        if committee_info:
            officer['committee_membership'] = '; '.join(set(committee_info))
        
        # Only return if we found at least a name
        return officer if officer.get('name') else None

    def parse_officer_table(self, table) -> List[Dict]:
        """Parse officer data from table format."""
        officers = []
        rows = table.find_all('tr')
        
        for row in rows[1:]:  # Skip header row
            cells = row.find_all(['td', 'th'])
            if len(cells) >= 2:
                officer = {
                    'name': cells[0].get_text().strip(),
                    'position': cells[1].get_text().strip() if len(cells) > 1 else '',
                    'description': cells[2].get_text().strip() if len(cells) > 2 else '',
                    'committee_membership': cells[3].get_text().strip() if len(cells) > 3 else ''
                }
                if officer['name']:
                    officers.append(officer)
        
        return officers

    def create_docx_report(self, officers: List[Dict], filename: str = "amazon_officers_directors.docx"):
        """Create a DOCX report with the officer data."""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Amazon.com Inc. - Officers and Directors', 0)
        
        # Add timestamp
        doc.add_paragraph(f'Data extracted on: {time.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Source: {self.base_url}')
        doc.add_paragraph('=' * 80)
        
        if not officers:
            doc.add_paragraph('No officer data was successfully extracted from the webpage.')
            doc.add_paragraph('This may be due to:')
            doc.add_paragraph('• Website access restrictions (403 Forbidden)')
            doc.add_paragraph('• Changes in website structure')
            doc.add_paragraph('• Anti-scraping measures')
        else:
            doc.add_paragraph(f'Total officers and directors found: {len(officers)}')
            doc.add_paragraph()
            
            for i, officer in enumerate(officers, 1):
                # Add officer section
                heading = doc.add_heading(f"{i}. {officer.get('name', 'Unknown')}", level=1)
                
                if officer.get('position'):
                    doc.add_paragraph(f"Position: {officer['position']}")
                
                if officer.get('description'):
                    doc.add_paragraph("Biography:")
                    bio_para = doc.add_paragraph(officer['description'])
                    bio_para.style = 'Quote'
                
                if officer.get('committee_membership'):
                    doc.add_paragraph(f"Committee Membership: {officer['committee_membership']}")
                
                # Add separator
                doc.add_paragraph('-' * 60)
        
        # Save document
        doc.save(filename)
        print(f"✓ Report saved as: {filename}")
        return filename

    def scrape_and_export(self):
        """Main method to scrape data and create DOCX export."""
        print("Starting Amazon Officers & Directors scraper...")
        print(f"Target URL: {self.base_url}")
        
        # Fetch the page
        soup = self.fetch_page()
        
        if not soup:
            print("✗ Failed to fetch the webpage")
            # Create a report indicating the failure
            self.create_docx_report([], "amazon_officers_directors_error.docx")
            return False
        
        print("✓ Page fetched successfully")
        
        # Extract officer data
        print("Extracting officer and director data...")
        self.officers_data = self.extract_officer_data(soup)
        
        if self.officers_data:
            print(f"✓ Successfully extracted {len(self.officers_data)} officer/director records")
            for officer in self.officers_data:
                print(f"  - {officer.get('name', 'N/A')}: {officer.get('position', 'N/A')}")
        else:
            print("⚠ No officer data could be extracted")
            print("The website may have anti-scraping measures or a different structure than expected")
        
        # Create DOCX report
        print("Creating DOCX report...")
        filename = self.create_docx_report(self.officers_data)
        
        print("\n" + "="*50)
        print("SCRAPING COMPLETE")
        print("="*50)
        print(f"Records extracted: {len(self.officers_data)}")
        print(f"Output file: {filename}")
        
        return True


def main():
    """Main execution function."""
    scraper = AmazonOfficersScraper()
    success = scraper.scrape_and_export()
    
    if success:
        print("\n✓ Scraping completed successfully!")
    else:
        print("\n✗ Scraping encountered issues. Check the error report.")


if __name__ == "__main__":
    main()